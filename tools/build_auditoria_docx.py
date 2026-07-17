from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "audit" / "Auditoria-Profesional-Completa-GamificApp.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"
LIGHT_BLUE_FILL = "E8EEF5"
RISK_RED = RGBColor(155, 28, 28)
OK_GREEN = RGBColor(34, 114, 69)
WARN_GOLD = RGBColor(122, 90, 0)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DEE7"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def no_split_table_rows(table):
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("GamificApp - Auditoria tecnica profesional")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Auditoria Profesional Completa")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("GamificApp - Plataforma Web de Gamificacion Educativa")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE

    meta = [
        ("Proyecto", "GamificApp"),
        ("Contexto", "Tesis universitaria - Ingenieria en Tecnologias de la Informacion"),
        ("Auditoria", "Arquitectura, frontend, backend, base de datos, seguridad, IA, gamificacion y UX/UI"),
        ("Fecha", "14 de julio de 2026"),
    ]
    add_kv_table(doc, meta, widths=(2200, 7160))
    add_callout(
        doc,
        "Conclusion ejecutiva",
        "GamificApp es un MVP academico avanzado y funcional. Puede presentarse ante un tribunal como proyecto de tesis si se corrigen los errores de lint, se documentan limitaciones y se prepara una demo controlada. No debe presentarse como producto listo para produccion sin pruebas, refactorizacion y endurecimiento de seguridad.",
    )


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, "CBD5E1")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE_FILL)
    set_cell_width(cell, 9360)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = DARK_BLUE
    r.font.size = Pt(10)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_kv_table(doc, rows, widths=(2600, 6760)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_width(cells[0], widths[0])
        set_cell_width(cells[1], widths[1])
        set_cell_shading(cells[0], GRAY_FILL)
        set_cell_text(cells[0], key, bold=True, color=DARK_BLUE)
        set_cell_text(cells[1], value)
    no_split_table_rows(table)
    doc.add_paragraph()


def add_matrix(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, GRAY_FILL)
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
        if widths:
            set_cell_width(cell, widths[i])
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                set_cell_width(cells[i], widths[i])
            if i == 1 and isinstance(value, str):
                if value.startswith("Cumple") or value.startswith("Alto"):
                    cells[i].paragraphs[0].runs[0].font.color.rgb = OK_GREEN
                elif value.startswith("Parcial") or value.startswith("Medio"):
                    cells[i].paragraphs[0].runs[0].font.color.rgb = WARN_GOLD
                elif value.startswith("No") or value.startswith("Bajo"):
                    cells[i].paragraphs[0].runs[0].font.color.rgb = RISK_RED
    no_split_table_rows(table)
    doc.add_paragraph()


def add_section(doc, title, good=None, improve=None, wrong=None):
    doc.add_heading(title, level=1)
    if good:
        doc.add_heading("Lo bueno", level=2)
        add_bullets(doc, good)
    if improve:
        doc.add_heading("Lo mejorable", level=2)
        add_bullets(doc, improve)
    if wrong:
        doc.add_heading("Incorrecto o riesgoso", level=2)
        add_bullets(doc, wrong)


def build():
    doc = Document()
    style_document(doc)
    add_footer(doc)
    add_title(doc)

    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "GamificApp es una plataforma web de gamificacion educativa dirigida a estudiantes de Educacion Basica Elemental. El proyecto incluye tres roles principales: administrador, docente y estudiante. Integra actividades gamificadas, XP, ranking, misiones, recursos educativos, juegos y generacion asistida por inteligencia artificial."
    )
    doc.add_paragraph(
        "La auditoria evidencia un proyecto funcional y ambicioso, con una base tecnica razonable para tesis. Sus principales riesgos son deuda tecnica acumulada en archivos monoliticos, ausencia de pruebas automatizadas, lint fallando, almacenamiento de archivos en base64 dentro de MySQL y una arquitectura de IA aun acoplada a Gemini."
    )
    add_kv_table(
        doc,
        [
            ("Build frontend", "Cumple: npm.cmd run build finalizo correctamente."),
            ("Lint", "No cumple: npm.cmd run lint reporto 29 errores y 3 warnings."),
            ("Bundle principal", "Advertencia: 1.47 MB minificado; requiere code splitting."),
            ("Backend", "Funcional: Express, JWT, MySQL, rutas por dominio y permisos."),
            ("Base de datos", "Razonable: relacional con FKs, pero con JSON polimorfico y archivos base64."),
        ],
    )

    add_section(
        doc,
        "2. Arquitectura",
        good=[
            "Separacion fisica entre frontend, backend y base de datos: src, server y database.",
            "Rutas backend por dominio funcional: auth, admin, docente, retos, misiones, ranking e IA.",
            "Servicios frontend separados para consumo de API.",
            "Middleware centralizado para autenticacion JWT, roles y permisos administrativos.",
            "Migraciones SQL y documentacion tecnica en docs.",
        ],
        improve=[
            "La navegacion interna se concentra en una ruta unica /dashboard con estado local, lo que limita deep linking, historial del navegador y mantenibilidad.",
            "No existe Clean Architecture estricta: controladores, reglas de negocio, validaciones y SQL conviven dentro de routers.",
            "Falta capa de servicios/repositorios backend para separar acceso a datos y reglas de dominio.",
            "El sistema de juegos esta parcialmente modularizado, pero aun requiere registros manuales paralelos.",
        ],
        wrong=[
            "Archivos criticos exceden tamanos saludables: dashboard.jsx supera 1000 lineas, admin.js supera 1000 lineas y dashboard.css se acerca a 2000 lineas.",
            "No se observaron pruebas automatizadas como parte del flujo de calidad.",
            "Persistir archivos como data_url base64 en MySQL compromete escalabilidad.",
        ],
    )

    add_section(
        doc,
        "3. Frontend",
        good=[
            "Uso de React, Vite, React Router y MUI.",
            "Componentes reutilizables para dashboard, archivos, juegos y quiz.",
            "Experiencias diferenciadas para administrador, docente y estudiante.",
            "Hay uso de hooks y memoizacion en varias pantallas.",
            "Los juegos tienen componentes propios y un registro central de frontend.",
        ],
        improve=[
            "Implementar React.lazy y code splitting por rol, juegos, PDF y Office.",
            "Reemplazar alerts por toasts o modales accesibles.",
            "Mejorar accesibilidad: foco, ARIA, navegacion por teclado, contraste y etiquetas.",
            "Separar utilidades exportadas desde archivos .jsx para resolver errores de Fast Refresh.",
            "Reducir uso de CSS global compartido entre roles.",
        ],
        wrong=[
            "El lint falla por patrones de hooks, Fast Refresh y asignaciones inutiles.",
            "La UX del estudiante hereda patrones de dashboard administrativo, poco ideal para ninos de 2do a 4to grado.",
            "El bundle principal es demasiado grande para una experiencia educativa ligera.",
        ],
    )

    add_section(
        doc,
        "4. Backend",
        good=[
            "JWT con secreto obligatorio, sin fallback inseguro.",
            "Passwords y PINs hasheados con bcrypt.",
            "Rate limiting basico para rutas publicas de autenticacion.",
            "Bloqueo por intentos fallidos.",
            "Permisos administrativos revalidados contra base de datos.",
            "Consultas SQL mayoritariamente parametrizadas.",
        ],
        improve=[
            "Agregar validacion formal de payloads con Zod, Joi o Yup.",
            "Dividir rutas grandes en controladores, servicios y repositorios.",
            "Agregar OpenAPI/Swagger.",
            "Agregar logging estructurado y trazabilidad de errores.",
            "Reemplazar rate limiting en memoria por Redis o mecanismo persistente si escala a multiples instancias.",
        ],
        wrong=[
            "El limite JSON de 25 MB para materiales base64 es riesgoso.",
            "No hay pruebas automatizadas de endpoints criticos.",
            "No hay sanitizacion formal centralizada para todos los campos de entrada.",
        ],
    )

    doc.add_heading("5. Base de datos", level=1)
    doc.add_paragraph(
        "El modelo relacional cubre las entidades centrales de la plataforma: usuarios, estudiantes, cursos, materias, retos, progreso, materiales, docentes, invitaciones, misiones y auditoria. La estructura es suficiente para una tesis funcional, pero necesita mayor normalizacion y granularidad si se desea analitica avanzada."
    )
    add_matrix(
        doc,
        ["Aspecto", "Evaluacion", "Observacion"],
        [
            ("Normalizacion", "Parcial", "curso y curso_id duplican informacion por compatibilidad."),
            ("Relaciones", "Cumple", "Existen FKs en entidades principales."),
            ("Indices", "Parcial", "Hay indices utiles, pero faltan para filtros frecuentes por estado, curso, docente y fechas."),
            ("Escalabilidad", "Parcial", "JSON polimorfico y archivos base64 limitan crecimiento."),
            ("Analitica", "Parcial", "No hay respuestas por pregunta ni intentos historicos detallados."),
        ],
        widths=[2200, 1600, 5560],
    )

    doc.add_heading("6. Seguridad", level=1)
    add_bullets(
        doc,
        [
            "Fortalezas: JWT, bcrypt, CORS configurable, cabeceras basicas, HSTS, rate limiting y permisos por modulo.",
            "Riesgo principal: JWT almacenado en localStorage; ante XSS el token puede ser extraido.",
            "Falta Helmet completo, validacion centralizada y politicas de subida de archivos mas estrictas.",
            "La subida de archivos requiere validacion MIME/backend, antivirus o sandbox si se lleva a produccion.",
            "Recomendacion: migrar a cookies HttpOnly Secure SameSite o endurecer CSP y sanitizacion si se conserva Bearer token.",
        ]
    )

    doc.add_heading("7. Rendimiento", level=1)
    add_bullets(
        doc,
        [
            "El build de produccion compila, pero Vite advierte chunks mayores a 500 KB.",
            "El bundle principal pesa 1.47 MB minificado y el PDF worker 1.29 MB.",
            "Debe implementarse lazy loading por rol y por modulos pesados.",
            "Los materiales base64 pueden inflar respuestas y memoria del servidor.",
            "El polling automatico puede escalar mal si aumentan usuarios concurrentes.",
        ]
    )

    doc.add_heading("8. Inteligencia artificial", level=1)
    doc.add_paragraph(
        "El cliente de IA esta implementado alrededor de Gemini. Es positivo que la API key viva en backend, que se solicite JSON estructurado y que existan reintentos ante saturacion. Sin embargo, el diseno esta acoplado a un unico proveedor."
    )
    add_callout(
        doc,
        "Recomendacion IA",
        "Implementar una arquitectura Provider para soportar OpenAI, Gemini, Claude y Ollama mediante un contrato comun: generateJSON, generateText, healthCheck y estimateCost.",
    )

    doc.add_heading("9. Gamificacion", level=1)
    add_matrix(
        doc,
        ["Elemento", "Estado", "Comentario"],
        [
            ("XP", "Cumple", "Existe acumulacion de XP y progreso por reto."),
            ("Niveles", "Parcial", "Se observa calculo/uso de nivel, pero no motor configurable robusto."),
            ("Insignias", "Parcial", "Existen logros visuales, pero falta persistencia formal completa."),
            ("Misiones", "Cumple", "Sistema de misiones con semillas y progreso."),
            ("Leaderboard", "Cumple", "Ranking disponible."),
            ("Recompensas", "Parcial", "No hay economia de recompensas configurable."),
            ("Retroalimentacion", "Parcial", "Existe feedback, pero falta motor por rangos de puntaje."),
        ],
        widths=[2200, 1600, 5560],
    )

    doc.add_heading("10. Juegos y arquitectura plugin", level=1)
    doc.add_paragraph(
        "El proyecto ya posee una base para modularizacion de juegos mediante un registro frontend. Los tipos actuales incluyen clasificador, memorama, linea de tiempo y completar espacios. Para escalar, cada juego deberia convertirse en un modulo independiente con manifest, Player, Editor, validator, aiSchema y politica de XP."
    )
    add_numbered(
        doc,
        [
            "Crear carpeta games/<tipo> por juego.",
            "Definir manifest.js con type, label, version, player, editor, validateConfig, summarize y xpPolicy.",
            "Cargar Player y Editor con dynamic import.",
            "Centralizar validadores backend y frontend desde un contrato compartido.",
            "Permitir activar/desactivar juegos desde administracion.",
        ]
    )

    doc.add_heading("11. Comparacion con la tesis", level=1)
    add_matrix(
        doc,
        ["Modulo", "Estado", "Detalle"],
        [
            ("Administrador", "Cumple", "Panel y permisos administrativos implementados."),
            ("Docente", "Cumple", "Panel docente, actividades, recursos y estudiantes."),
            ("Estudiante", "Cumple", "Dashboard, juegos, quiz, misiones y progreso."),
            ("Materias", "Cumple", "Gestion y catalogo de materias."),
            ("Actividades", "Cumple", "Retos/actividades por tipo."),
            ("Banco de preguntas", "Parcial", "Preguntas viven dentro de configuracion de retos, no como banco independiente."),
            ("Evidencias", "No implementado", "No existe modulo formal de evidencias."),
            ("Repositorio", "Parcial", "Materiales funcionan como repositorio simple."),
            ("Asistencia", "No implementado", "No existe modulo de asistencia."),
            ("Recursos", "Cumple", "Carga y consulta de materiales."),
            ("Reportes", "Parcial", "Hay dashboards/ranking/libro, faltan reportes exportables completos."),
            ("Dashboard", "Cumple", "Paneles por rol."),
            ("IA", "Cumple", "Generacion con Gemini."),
            ("Gamificacion", "Parcial", "Base solida, falta motor configurable completo."),
        ],
        widths=[2400, 1700, 5260],
    )

    doc.add_heading("12. Observaciones del revisor", level=1)
    add_matrix(
        doc,
        ["Punto", "Estado", "Observacion"],
        [
            ("Funcionamiento local", "Cumple", "Frontend compila; backend requiere MySQL y variables de entorno."),
            ("Carga masiva mediante Excel", "No implementado", "No se detecto flujo de importacion masiva."),
            ("Generacion automatica de codigos", "Cumple", "Codigos de invitacion y emergencia implementados."),
            ("Seleccion por nombre + codigo", "Parcial", "Registro/emergencia usan nombre + codigo; login usa nombre + PIN."),
            ("Administracion de juegos", "Parcial", "Se administran retos/actividades, no catalogo formal de juegos."),
            ("Aleatorizacion de respuestas", "Parcial", "Juegos mezclan elementos; quiz mantiene A-D."),
            ("Banco de preguntas", "Parcial", "No existe entidad independiente."),
            ("Retroalimentacion por puntaje", "Parcial", "Hay justificacion y progreso; falta regla configurable por rangos."),
            ("Integracion con IA", "Cumple", "Generacion de actividades con Gemini."),
        ],
        widths=[2700, 1700, 4960],
    )

    doc.add_heading("13. Deuda tecnica priorizada", level=1)
    add_matrix(
        doc,
        ["Prioridad", "Deuda", "Impacto"],
        [
            ("Alta", "Corregir lint completo.", "Calidad y confianza de entrega."),
            ("Alta", "Dividir dashboard.jsx, AdminDashboard.jsx y admin.js.", "Mantenibilidad."),
            ("Alta", "Agregar pruebas backend/frontend.", "Regresiones y defensa tecnica."),
            ("Alta", "Validacion formal de requests.", "Seguridad y consistencia."),
            ("Alta", "Sacar archivos base64 de MySQL.", "Escalabilidad y rendimiento."),
            ("Media", "Arquitectura Provider para IA.", "Flexibilidad tecnologica."),
            ("Media", "Rutas anidadas por modulo.", "UX, navegacion y mantenibilidad."),
            ("Media", "Reportes exportables.", "Cumplimiento academico/operativo."),
            ("Baja", "Corregir mojibake en comentarios.", "Presentacion y legibilidad."),
            ("Baja", "Reemplazar alerts por componentes UI.", "Pulido UX."),
        ],
        widths=[1400, 4200, 3760],
    )

    doc.add_heading("14. Roadmap", level=1)
    add_matrix(
        doc,
        ["Version", "Objetivos principales", "Resultado esperado"],
        [
            ("1.1", "Corregir lint, dividir archivos grandes, lazy loading, validaciones backend y tests minimos.", "Entrega academica estabilizada."),
            ("1.2", "Banco de preguntas, carga masiva Excel, reportes, evidencias, insignias persistentes y almacenamiento externo.", "Plataforma escolar mas completa."),
            ("2.0", "IA multi-provider, motor gamificacion configurable, analitica por pregunta, asistencia, multi-institucion y plugins de juegos.", "Producto escalable."),
        ],
        widths=[1300, 5300, 2760],
    )

    doc.add_heading("15. Puntuacion final", level=1)
    add_matrix(
        doc,
        ["Area", "Puntuacion", "Justificacion"],
        [
            ("Arquitectura", "7/10", "Buena separacion fisica, pero sin capas limpias completas."),
            ("Backend", "7/10", "Funcional y seguro en bases, pero monolitico y sin validacion formal."),
            ("Frontend", "6.5/10", "Completo, pero pesado y con lint fallando."),
            ("Base de datos", "7/10", "Relacional razonable, limitada por JSON y base64."),
            ("Seguridad", "6.5/10", "Buenas bases, pero JWT en localStorage y upload mejorable."),
            ("Escalabilidad", "6/10", "Necesita almacenamiento externo, code splitting y mejor arquitectura."),
            ("Gamificacion", "7.5/10", "Base fuerte con XP, misiones y ranking."),
            ("UX/UI", "6.5/10", "Funcional, pero estudiante requiere experiencia mas infantil."),
            ("Calidad general", "7/10", "MVP academico avanzado, no producto final."),
        ],
        widths=[2400, 1500, 5460],
    )

    add_callout(
        doc,
        "Veredicto para tribunal",
        "Si. El proyecto puede presentarse ante un tribunal de Ingenieria como MVP academico avanzado, siempre que se corrijan errores visibles de calidad, se documente la deuda tecnica y se delimite claramente que no es una version productiva final.",
    )

    doc.add_heading("16. Evidencia revisada", level=1)
    add_bullets(
        doc,
        [
            "package.json y server/package.json para stack y dependencias.",
            "server/server.js, server/middleware/auth.js y rutas backend principales.",
            "database/gamificapp.sql y migraciones para modelo de datos.",
            "src/App.jsx, servicios frontend y componentes de juegos/quiz/dashboard.",
            "Ejecucion local de build y lint.",
        ]
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
